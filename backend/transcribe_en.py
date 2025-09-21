from typing import List, Dict, Tuple
import json
from pydub import AudioSegment
from faster_whisper import WhisperModel
from docx import Document
from docx.shared import RGBColor, Pt
from your_package_name import get_large_whisper

# ——— Helpers ———
def delete_paragraph(paragraph):
    """Hard-delete a paragraph from the doc body."""
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def add_turn(doc, speaker, text):
    """
    Add one dialogue line like:
    M: Hello there
    - whole line red if M
    - blank line (0.6 spacing) before each M entry
    """
    if speaker == "M":
        blank = doc.add_paragraph()
        blank.paragraph_format.line_spacing = 0.6

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)

    run = para.add_run(f"{speaker}: {text}")

    if speaker == "M":
        run.font.color.rgb = RGBColor(255, 0, 0)  # whole line red

# Load Whisper model
model = get_large_whisper()

# ——— Main function ———
def transcribe_en(
    audio_path: str,
    segments: List[Dict],
    template_path: str,
    output_docx: str,
    device: str = "cuda",
    beam_size: int = 5
) -> str:
    """
    Transcribe English audio with diarization and save directly to DOCX.

    Args:
        audio_path (str): Path to the input audio file (MP3/M4A/etc.)
        segments (List[Dict]): Diarization output (list of {start, end, speaker} dicts)
        template_path (str): Path to DOCX template (header/footer preserved)
        output_docx (str): Path where final DOCX will be saved
        device (str): Device for Whisper model
        beam_size (int): Beam size for Whisper

    Returns:
        str: Path to saved DOCX
    """
    # Load audio
    audio = AudioSegment.from_file(audio_path)
    audio = audio.set_channels(1).set_frame_rate(16000)

    # Transcribe each diarized segment
    turns: List[Tuple[str, str]] = []
    for idx, seg in enumerate(segments, start=1):
        start_ms = int(seg["start"] * 1000)
        end_ms = int(seg["end"] * 1000)
        speaker = seg["speaker"]

        snippet = audio[start_ms:end_ms]
        snippet_path = f"segment_{idx:03d}.wav"
        snippet.export(snippet_path, format="wav")

        w_segments, _ = model.transcribe(
            snippet_path,
            language="en",
            beam_size=beam_size,
        )
        text = " ".join(s.text.strip() for s in w_segments).strip()
        if text:
            turns.append((speaker, text))

    # Build the DOCX
    doc = Document(template_path)
    for p in list(doc.paragraphs):
        delete_paragraph(p)

    for speaker, text in turns:
        add_turn(doc, speaker, text)

    doc.save(output_docx)
    print(f"📄 Transcription saved to: {output_docx} | total turns: {len(turns)}")
    return output_docx
