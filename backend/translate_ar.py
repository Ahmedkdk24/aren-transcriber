from typing import List, Tuple
import re
from docx import Document
from docx.shared import RGBColor, Pt
from transformers import AutoModelForCausalLM, AutoTokenizer, pipeline, BitsAndBytesConfig
import os
import pickle
from your_package_name import get_text_gen_pipeline


# --- Global config ---
CHUNK_SIZE_WORDS = 300   # smaller chunks to ensure full translation fits

# --- Load model once globally ---
pipe = get_text_gen_pipeline()

# --- Helpers ---
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def add_turn(doc, speaker, text):
    """Add formatted dialogue line to doc."""
    if speaker == "M":
        blank = doc.add_paragraph()
        blank.paragraph_format.line_spacing = 0.6

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(f"{speaker}: {text}")
    run.font.size = Pt(12)

    if speaker == "M":
        run.font.color.rgb = RGBColor(255, 0, 0)

def chunk_turns(turns, max_words=CHUNK_SIZE_WORDS):
    """Group turns into word-limited chunks, ending only after a speaker label."""
    chunks, current_chunk, current_len = [], [], 0
    for speaker, text in turns:
        words = text.split()
        current_chunk.append((speaker, text))
        current_len += len(words)
        if current_len >= max_words:
            chunks.append(current_chunk)
            current_chunk, current_len = [], 0
    if current_chunk:
        chunks.append(current_chunk)
    return chunks

def llm_translate(dialogue: str, context_summary: str = "") -> str:
    """Translate dialogue chunk using ALLaM model."""
    prompt = f"""
You are translating a conversation from Arabic to English.
Preserve speaker labels ("M:" and "R:").
Translate faithfully and naturally into English.
Do not include Arabic in the output.

Now translate this part:
{dialogue}
    """.strip()

    messages = [{"role": "user", "content": prompt}]
    response = pipe(messages, max_new_tokens=1024, do_sample=False, temperature=0.0)

    # Some HF pipelines return plain string, some return dicts with "generated_text"
    if isinstance(response[0], dict) and "generated_text" in response[0]:
        return response[0]["generated_text"]
    return str(response[0])

def validate_translation(input_turns, final_turns):
    """Compare word counts between source and translated transcripts."""
    input_word_count = sum(len(text.split()) for _, text in input_turns)
    output_word_count = sum(len(text.split()) for _, text in final_turns)

    input_turns_count = len(input_turns)
    output_turns_count = len(final_turns)

    print("\n📊 Translation Validation Report")
    print(f"   Input  - {input_turns_count} turns, {input_word_count} words")
    print(f"   Output - {output_turns_count} turns, {output_word_count} words")

    if input_word_count > 0:
        coverage = (output_word_count / input_word_count) * 100
        print(f"   ✅ Coverage: {coverage:.2f}% of original words preserved in translation")

        if coverage < 90:
            print("   ⚠️ Warning: Significant word loss detected. Check raw chunk outputs.")
    else:
        print("   ⚠️ Input word count is zero. Something went wrong with the input data.")


# --- Main Function ---
def translate_ar(
    turns: List[Tuple[str, str]],
    template_path: str,
    output_docx: str,
    resume_progress: bool = False,
    progress_path: str = "/tmp/translation_progress.pkl"
) -> List[Tuple[str, str]]:
    """
    Translate Arabic transcript turns into English and save to DOCX.

    Args:
        turns (List[Tuple[str, str]]): Transcript as (speaker, text).
        template_path (str): Path to DOCX template.
        output_docx (str): Where to save final translated DOCX.
        resume_progress (bool): Resume from saved progress if True.
        progress_path (str): Path to save progress pickle.

    Returns:
        List[Tuple[str, str]]: Translated turns.
    """

    final_turns = []
    context_summary = ""
    start_chunk = 0

    # Step 1: Split into chunks
    chunks = chunk_turns(turns)
    print(f"🔹 Split transcript into {len(chunks)} chunks")

    # Resume logic
    if resume_progress and os.path.exists(progress_path):
        with open(progress_path, "rb") as f:
            progress = pickle.load(f)
        final_turns = progress.get("final_turns", [])
        context_summary = progress.get("context_summary", "")
        start_chunk = progress.get("last_chunk", 0) + 1
        print(f"⏩ Resuming from chunk {start_chunk+1}")

    # Step 2: Translate sequentially
    for idx, chunk in enumerate(chunks[start_chunk:], start=start_chunk):
        dialogue = "\n".join([f"{sp}: {txt}" for sp, txt in chunk])

        try:
            translated_text = llm_translate(dialogue, context_summary)
        except Exception as e:
            print(f"❌ Error during translation of chunk {idx+1}: {e}")
            with open(progress_path, "wb") as f:
                pickle.dump({
                    "final_turns": final_turns,
                    "context_summary": context_summary,
                    "last_chunk": idx - 1
                }, f)
            raise

        # Parse translated lines
        translated_chunk = []
        last_line, repeat_count = None, 0
        for line in translated_text.splitlines():
            if line == last_line:
                repeat_count += 1
                if repeat_count > 2:
                    continue
            else:
                repeat_count = 0
            last_line = line

            if ":" in line:
                sp, txt = line.split(":", 1)
                translated_chunk.append((sp.strip(), txt.strip()))
            else:
                if translated_chunk:
                    translated_chunk[-1] = (
                        translated_chunk[-1][0],
                        translated_chunk[-1][1] + " " + line.strip()
                    )

        final_turns.extend(translated_chunk)
        print(f"✅ Translated chunk {idx+1}/{len(chunks)}")

        # Save progress
        with open(progress_path, "wb") as f:
            pickle.dump({
                "final_turns": final_turns,
                "context_summary": context_summary,
                "last_chunk": idx
            }, f)

    # Step 3: Build DOCX
    doc = Document(template_path)
    for p in list(doc.paragraphs):
        delete_paragraph(p)

    for speaker, text in final_turns:
        add_turn(doc, speaker, text)

    doc.save(output_docx)
    print(f"📄 English transcription saved to: {output_docx}")

    validate_translation(turns, final_turns)
    return final_turns
